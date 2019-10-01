from docx import Document

document = Document('./demo.docx')

# 将word文档的内容一行一行的读取
for paragraph in document.paragraphs:
    print(paragraph.text)


for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
